"""
Module xuất báo cáo Word
Xuất 14 chỉ số, PD, biểu đồ và phân tích của Gemini ra file Word
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List
import matplotlib.pyplot as plt
import io
import os
from datetime import datetime


class ReportGenerator:
    """Class để tạo báo cáo Word"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()

    def setup_styles(self):
        """Thiết lập styles cho document"""
        # Thiết lập font mặc định
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

    def add_header(self):
        """Thêm header cho báo cáo"""
        # Title
        title = self.doc.add_heading('BÁO CÁO ĐÁNH GIÁ RỦI RO TÍN DỤNG DOANH NGHIỆP', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(0, 166, 81)  # Màu xanh Agribank
        title_run.font.bold = True

        # Subtitle
        subtitle = self.doc.add_paragraph('HỆ THỐNG AI STACKING CLASSIFIER')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(14)
        subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

        # Ngày tạo báo cáo
        date_para = self.doc.add_paragraph(f'Ngày tạo báo cáo: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.runs[0]
        date_run.font.size = Pt(11)
        date_run.font.italic = True

        self.doc.add_paragraph()  # Spacing

    def add_section_title(self, title: str):
        """Thêm tiêu đề section"""
        heading = self.doc.add_heading(title, level=1)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(0, 166, 81)

    def add_prediction_results(self, prediction: Dict[str, Any]):
        """Thêm phần kết quả dự báo PD"""
        self.add_section_title('I. KẾT QUẢ DỰ BÁO XÁC SUẤT VỠ NỢ (PD)')

        # Tạo bảng
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Mô hình'
        header_cells[1].text = 'Xác suất Vỡ nợ (PD)'

        # Dữ liệu
        models = [
            ('Stacking Model (Kết quả chính)', prediction.get('pd_stacking', 0) * 100),
            ('Logistic Regression', prediction.get('pd_logistic', 0) * 100),
            ('Random Forest', prediction.get('pd_random_forest', 0) * 100),
            ('XGBoost', prediction.get('pd_xgboost', 0) * 100)
        ]

        for i, (model_name, pd_value) in enumerate(models, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = model_name
            row_cells[1].text = f'{pd_value:.2f}%'

        # Kết luận
        self.doc.add_paragraph()
        conclusion = self.doc.add_paragraph()
        conclusion.add_run('Kết luận: ').bold = True
        conclusion.add_run(prediction.get('prediction_label', 'N/A'))

        self.doc.add_paragraph()

    def add_14_indicators(self, indicators: List[Dict[str, Any]]):
        """Thêm phần 14 chỉ số tài chính"""
        self.add_section_title('II. 14 CHỈ SỐ TÀI CHÍNH')

        # Tạo bảng
        table = self.doc.add_table(rows=len(indicators) + 1, cols=3)
        table.style = 'Light Grid Accent 1'

        # Header
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Mã chỉ số'
        header_cells[1].text = 'Tên chỉ số'
        header_cells[2].text = 'Giá trị'

        # Dữ liệu
        for i, indicator in enumerate(indicators, 1):
            row_cells = table.rows[i].cells
            row_cells[0].text = indicator['code']
            row_cells[1].text = indicator['name']
            row_cells[2].text = f"{indicator['value']:.6f}"

        self.doc.add_paragraph()

    def add_chart(self, prediction: Dict[str, Any], indicators_dict: Dict[str, float]):
        """Thêm biểu đồ vào báo cáo"""
        self.add_section_title('III. BIỂU ĐỒ PHÂN TÍCH')

        # Biểu đồ 1: So sánh PD từ 4 models
        self.doc.add_heading('3.1. So sánh Xác suất Vỡ nợ (PD) từ 4 Models', level=2)

        fig1, ax1 = plt.subplots(figsize=(10, 6))
        models = ['Stacking', 'Logistic', 'Random Forest', 'XGBoost']
        pd_values = [
            prediction.get('pd_stacking', 0) * 100,
            prediction.get('pd_logistic', 0) * 100,
            prediction.get('pd_random_forest', 0) * 100,
            prediction.get('pd_xgboost', 0) * 100
        ]

        colors = ['#00a651', '#4CAF50', '#8BC34A', '#CDDC39']
        bars = ax1.bar(models, pd_values, color=colors)
        ax1.set_ylabel('Xác suất Vỡ nợ (%)', fontsize=12)
        ax1.set_title('So sánh PD từ 4 Models', fontsize=14, fontweight='bold')
        ax1.set_ylim(0, 100)

        # Thêm giá trị trên các cột
        for bar, value in zip(bars, pd_values):
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height,
                    f'{value:.2f}%',
                    ha='center', va='bottom', fontsize=10)

        # Lưu biểu đồ
        chart1_path = 'chart_pd_comparison.png'
        plt.tight_layout()
        plt.savefig(chart1_path, dpi=300, bbox_inches='tight')
        plt.close()

        # Thêm vào document
        self.doc.add_picture(chart1_path, width=Inches(6))
        os.remove(chart1_path)

        self.doc.add_paragraph()

        # Biểu đồ 2: 14 chỉ số tài chính
        self.doc.add_heading('3.2. Phân tích 14 Chỉ số Tài chính', level=2)

        fig2, (ax2_1, ax2_2) = plt.subplots(1, 2, figsize=(14, 6))

        # Biểu đồ 2a: Nhóm chỉ số sinh lời và đòn bẩy (X1-X6)
        indicators_group1 = ['X_1', 'X_2', 'X_3', 'X_4', 'X_5', 'X_6']
        values_group1 = [indicators_dict.get(key, 0) for key in indicators_group1]
        labels_group1 = [f'{key}' for key in indicators_group1]

        ax2_1.barh(labels_group1, values_group1, color='#FFB6C1')
        ax2_1.set_xlabel('Giá trị', fontsize=10)
        ax2_1.set_title('Nhóm 1: Sinh lời & Đòn bẩy (X1-X6)', fontsize=11, fontweight='bold')
        ax2_1.grid(axis='x', alpha=0.3)

        # Biểu đồ 2b: Nhóm thanh toán và hiệu quả (X7-X14)
        indicators_group2 = ['X_7', 'X_8', 'X_9', 'X_10', 'X_11', 'X_12', 'X_13', 'X_14']
        values_group2 = [indicators_dict.get(key, 0) for key in indicators_group2]
        labels_group2 = [f'{key}' for key in indicators_group2]

        ax2_2.barh(labels_group2, values_group2, color='#ADD8E6')
        ax2_2.set_xlabel('Giá trị', fontsize=10)
        ax2_2.set_title('Nhóm 2: Thanh toán & Hiệu quả (X7-X14)', fontsize=11, fontweight='bold')
        ax2_2.grid(axis='x', alpha=0.3)

        # Lưu biểu đồ
        chart2_path = 'chart_indicators.png'
        plt.tight_layout()
        plt.savefig(chart2_path, dpi=300, bbox_inches='tight')
        plt.close()

        # Thêm vào document
        self.doc.add_picture(chart2_path, width=Inches(6.5))
        os.remove(chart2_path)

        self.doc.add_paragraph()

    def add_gemini_analysis(self, analysis: str):
        """Thêm phần phân tích của Gemini"""
        self.add_section_title('IV. PHÂN TÍCH CHUYÊN SÂU VÀ KHUYẾN NGHỊ')

        # Thêm nội dung phân tích
        paragraphs = analysis.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = self.doc.add_paragraph(para_text.strip())
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        self.doc.add_paragraph()

    def add_footer(self):
        """Thêm footer cho báo cáo"""
        self.doc.add_paragraph()
        self.doc.add_paragraph('_' * 80)

        footer = self.doc.add_paragraph()
        footer.add_run('Báo cáo được tạo tự động bởi ').italic = True
        footer.add_run('Hệ thống Đánh giá Rủi ro Tín dụng - Agribank').bold = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        disclaimer = self.doc.add_paragraph(
            'Lưu ý: Báo cáo này chỉ mang tính chất tham khảo. '
            'Quyết định cho vay cần được xem xét bởi các chuyên gia tín dụng và tuân thủ quy trình nội bộ của ngân hàng.'
        )
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        disclaimer_run = disclaimer.runs[0]
        disclaimer_run.font.size = Pt(10)
        disclaimer_run.font.italic = True
        disclaimer_run.font.color.rgb = RGBColor(150, 150, 150)

    def generate_report(self, data: Dict[str, Any], output_path: str = 'bao_cao_tin_dung.docx') -> str:
        """
        Tạo báo cáo hoàn chỉnh

        Args:
            data: Dữ liệu bao gồm prediction, indicators, và analysis
            output_path: Đường dẫn file output

        Returns:
            Đường dẫn file báo cáo
        """
        try:
            # Lấy dữ liệu
            prediction = data.get('prediction', {})
            indicators = data.get('indicators', [])
            indicators_dict = data.get('indicators_dict', {})
            analysis = data.get('analysis', 'Không có phân tích')

            # Tạo báo cáo
            self.add_header()
            self.add_prediction_results(prediction)
            self.add_14_indicators(indicators)
            self.add_chart(prediction, indicators_dict)
            self.add_gemini_analysis(analysis)
            self.add_footer()

            # Lưu file
            self.doc.save(output_path)
            return output_path

        except Exception as e:
            raise ValueError(f"Lỗi khi tạo báo cáo: {str(e)}")

    def generate_survival_report(self, data: Dict[str, Any], output_path: str = 'bao_cao_survival_analysis.docx') -> str:
        """
        Tạo báo cáo Survival Analysis

        Args:
            data: Dữ liệu bao gồm:
                - indicators: 14 chỉ số tài chính
                - median_time_to_default: Median time
                - survival_probabilities: Xác suất sống sót tại 6/12/24 tháng
                - risk_classification: Phân loại rủi ro
                - hazard_ratios: Top hazard ratios
                - survival_curve: Timeline và probabilities
                - gemini_analysis: Phân tích từ Gemini
                - warning: Cảnh báo (nếu có)
            output_path: Đường dẫn file output

        Returns:
            Đường dẫn file báo cáo
        """
        try:
            # Tạo document mới
            self.doc = Document()
            self.setup_styles()

            # Lấy dữ liệu
            indicators = data.get('indicators', {})
            median_time = data.get('median_time_to_default', 0)
            survival_probs = data.get('survival_probabilities', {})
            risk_info = data.get('risk_classification', {})
            hazard_ratios = data.get('hazard_ratios', [])
            survival_curve = data.get('survival_curve', {})
            gemini_analysis = data.get('gemini_analysis', 'Không có phân tích')
            warning = data.get('warning', None)

            # ================================
            # HEADER
            # ================================
            title = self.doc.add_heading('BÁO CÁO PHÂN TÍCH SỐNG SÓT & TIME-TO-DEFAULT', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.color.rgb = RGBColor(0, 166, 81)
            title_run.font.bold = True

            subtitle = self.doc.add_paragraph('HỆ THỐNG SURVIVAL ANALYSIS - COX PROPORTIONAL HAZARDS MODEL')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.color.rgb = RGBColor(100, 100, 100)

            date_para = self.doc.add_paragraph(f'Ngày tạo báo cáo: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.runs[0]
            date_run.font.size = Pt(11)
            date_run.font.italic = True

            self.doc.add_paragraph()

            # ================================
            # I. TỔNG QUAN KẾT QUẢ
            # ================================
            self.add_section_title('I. TỔNG QUAN KẾT QUẢ SURVIVAL ANALYSIS')

            # Median Time-to-Default
            median_para = self.doc.add_paragraph()
            median_para.add_run('Median Time-to-Default: ').bold = True
            median_para.add_run(f'{median_time:.1f} tháng')
            median_run = median_para.runs[1]
            median_run.font.size = Pt(14)
            median_run.font.color.rgb = RGBColor(220, 20, 60) if median_time < 12 else RGBColor(0, 166, 81)

            # Giải thích
            explanation = self.doc.add_paragraph(
                f'Doanh nghiệp có 50% xác suất vỡ nợ trong vòng {median_time:.1f} tháng tới.'
            )
            explanation_run = explanation.runs[0]
            explanation_run.font.italic = True

            # Phân loại rủi ro
            risk_para = self.doc.add_paragraph()
            risk_para.add_run('Mức độ Rủi ro: ').bold = True
            risk_para.add_run(f'{risk_info.get("level", "N/A")}')
            risk_para.add_run(f'\n{risk_info.get("description", "")}')

            self.doc.add_paragraph()

            # Cảnh báo nếu có
            if warning:
                warning_para = self.doc.add_paragraph()
                warning_para.add_run('⚠️ CẢNH BÁO: ').bold = True
                warning_para.add_run(warning.get('message', ''))
                warning_run = warning_para.runs[0]
                warning_run.font.color.rgb = RGBColor(220, 20, 60)
                warning_run.font.size = Pt(13)

                rec_para = self.doc.add_paragraph()
                rec_para.add_run('Khuyến nghị: ').bold = True
                rec_para.add_run(warning.get('recommendation', ''))

                self.doc.add_paragraph()

            # ================================
            # II. XÁC SUẤT SỐNG SÓT THEO THỜI GIAN
            # ================================
            self.add_section_title('II. XÁC SUẤT SỐNG SÓT THEO THỜI GIAN')

            # Tạo bảng xác suất
            prob_table = self.doc.add_table(rows=4, cols=3)
            prob_table.style = 'Light Grid Accent 1'

            # Header
            header_cells = prob_table.rows[0].cells
            header_cells[0].text = 'Thời điểm'
            header_cells[1].text = 'Xác suất Sống sót'
            header_cells[2].text = 'Xác suất Vỡ nợ'

            # Dữ liệu
            for i, time in enumerate([6, 12, 24], 1):
                surv_prob = survival_probs.get(time, 0)
                default_prob = 1 - surv_prob

                row_cells = prob_table.rows[i].cells
                row_cells[0].text = f'{time} tháng'
                row_cells[1].text = f'{surv_prob * 100:.2f}%'
                row_cells[2].text = f'{default_prob * 100:.2f}%'

            self.doc.add_paragraph()

            # ================================
            # III. 14 CHỈ SỐ TÀI CHÍNH
            # ================================
            self.add_section_title('III. 14 CHỈ SỐ TÀI CHÍNH DOANH NGHIỆP')

            indicator_names = {
                'X_1': 'Biên lợi nhuận gộp',
                'X_2': 'Biên lợi nhuận trước thuế',
                'X_3': 'ROA (Return on Assets)',
                'X_4': 'ROE (Return on Equity)',
                'X_5': 'Hệ số nợ trên tài sản',
                'X_6': 'Hệ số nợ trên VCSH',
                'X_7': 'Khả năng thanh toán hiện hành',
                'X_8': 'Khả năng thanh toán nhanh',
                'X_9': 'Khả năng trả lãi',
                'X_10': 'Khả năng trả nợ gốc',
                'X_11': 'Khả năng tạo tiền/VCSH',
                'X_12': 'Vòng quay hàng tồn kho',
                'X_13': 'Kỳ thu tiền bình quân',
                'X_14': 'Hiệu suất sử dụng tài sản'
            }

            # Tạo bảng chỉ số
            indicators_table = self.doc.add_table(rows=15, cols=3)
            indicators_table.style = 'Light Grid Accent 1'

            # Header
            header_cells = indicators_table.rows[0].cells
            header_cells[0].text = 'Mã'
            header_cells[1].text = 'Tên chỉ số'
            header_cells[2].text = 'Giá trị'

            # Dữ liệu
            for i, key in enumerate(sorted(indicators.keys()), 1):
                if key in indicator_names:
                    row_cells = indicators_table.rows[i].cells
                    row_cells[0].text = key
                    row_cells[1].text = indicator_names[key]
                    row_cells[2].text = f'{indicators[key]:.4f}'

            self.doc.add_paragraph()

            # ================================
            # IV. YẾU TỐ RỦI RO - RISK CONTRIBUTIONS hoặc HAZARD RATIOS
            # ================================

            # Ưu tiên hiển thị risk_contributions (individual) nếu có
            risk_contributions = data.get('risk_contributions', [])

            if risk_contributions:
                # Hiển thị INDIVIDUAL RISK CONTRIBUTIONS (cụ thể cho doanh nghiệp này)
                self.add_section_title('IV. YẾU TỐ RỦI RO - ĐÓNG GÓP VÀO RỦI RO CỦA DOANH NGHIỆP NÀY')

                # Giải thích
                explanation = self.doc.add_paragraph(
                    'Risk Contribution cho biết mỗi chỉ số tài chính đóng góp như thế nào vào rủi ro CỤ THỂ của doanh nghiệp này:'
                )
                explanation.add_run('\n• (+): Chỉ số này TĂNG rủi ro vỡ nợ cho doanh nghiệp')
                explanation.add_run('\n• (-): Chỉ số này GIẢM rủi ro vỡ nợ cho doanh nghiệp')
                explanation.add_run('\n• Giá trị càng lớn (dương hoặc âm) = Ảnh hưởng càng mạnh')

                self.doc.add_paragraph()

                # Tạo bảng risk contributions
                rc_table = self.doc.add_table(rows=len(risk_contributions) + 1, cols=5)
                rc_table.style = 'Light Grid Accent 1'

                # Header
                header_cells = rc_table.rows[0].cells
                header_cells[0].text = 'Chỉ số'
                header_cells[1].text = 'Giá trị DN'
                header_cells[2].text = 'So với TB'
                header_cells[3].text = 'Đóng góp'
                header_cells[4].text = 'Diễn giải'

                # Dữ liệu
                for i, rc in enumerate(risk_contributions, 1):
                    feature_name = rc.get('feature_name', 'N/A')
                    company_value = rc.get('company_value', 0)
                    comparison = rc.get('comparison', 'N/A')
                    contribution = rc.get('risk_contribution', 0)
                    interpretation = rc.get('interpretation', 'N/A')

                    row_cells = rc_table.rows[i].cells
                    row_cells[0].text = feature_name
                    row_cells[1].text = f'{company_value:.3f}'
                    row_cells[2].text = comparison
                    row_cells[3].text = f'{contribution:+.3f}'
                    row_cells[4].text = interpretation

            elif hazard_ratios:
                # Hiển thị HAZARD RATIOS (model-level, tổng quát)
                self.add_section_title('IV. HAZARD RATIOS - YẾU TỐ RỦI RO QUAN TRỌNG (Model-Level)')

                # Giải thích Hazard Ratio
                explanation = self.doc.add_paragraph(
                    'Hazard Ratio (HR) đo lường mức độ ảnh hưởng TRUNG BÌNH của từng chỉ số tài chính đến rủi ro vỡ nợ:'
                )
                explanation.add_run('\n• HR > 1: Chỉ số này làm TĂNG nguy cơ vỡ nợ (càng cao càng nguy hiểm)')
                explanation.add_run('\n• HR < 1: Chỉ số này làm GIẢM nguy cơ vỡ nợ (bảo vệ doanh nghiệp)')
                explanation.add_run('\n• HR = 1: Chỉ số không ảnh hưởng đến rủi ro')

                self.doc.add_paragraph()

                # Tạo bảng hazard ratios
                hr_table = self.doc.add_table(rows=len(hazard_ratios) + 1, cols=4)
                hr_table.style = 'Light Grid Accent 1'

                # Header
                header_cells = hr_table.rows[0].cells
                header_cells[0].text = 'Chỉ số'
                header_cells[1].text = 'Hazard Ratio'
                header_cells[2].text = 'Diễn giải'
                header_cells[3].text = 'Ý nghĩa'

                # Dữ liệu
                for i, hr in enumerate(hazard_ratios, 1):
                    feature_name = hr.get('feature_name', 'N/A')
                    ratio = hr.get('hazard_ratio', 1.0)
                    significance = hr.get('significance', 'N/A')

                    # Diễn giải
                    if ratio > 1:
                        interpretation = f'Tăng rủi ro {(ratio - 1) * 100:.1f}%'
                    elif ratio < 1:
                        interpretation = f'Giảm rủi ro {(1 - ratio) * 100:.1f}%'
                    else:
                        interpretation = 'Không ảnh hưởng'

                    row_cells = hr_table.rows[i].cells
                    row_cells[0].text = feature_name
                    row_cells[1].text = f'{ratio:.3f}'
                    row_cells[2].text = interpretation
                    row_cells[3].text = significance

            self.doc.add_paragraph()

            # ================================
            # V. PHÂN TÍCH GEMINI AI
            # ================================
            self.add_section_title('V. PHÂN TÍCH CHUYÊN SÂU TỪ GEMINI AI')

            # Thêm nội dung phân tích
            paragraphs = gemini_analysis.split('\n')
            for para_text in paragraphs:
                if para_text.strip():
                    para = self.doc.add_paragraph(para_text.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            self.doc.add_paragraph()

            # ================================
            # FOOTER
            # ================================
            self.doc.add_paragraph()
            self.doc.add_paragraph('_' * 80)

            footer = self.doc.add_paragraph()
            footer.add_run('Báo cáo được tạo tự động bởi ').italic = True
            footer.add_run('Hệ thống Survival Analysis - Agribank').bold = True
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

            disclaimer = self.doc.add_paragraph(
                'Lưu ý: Báo cáo này chỉ mang tính chất tham khảo. '
                'Quyết định cho vay cần được xem xét bởi các chuyên gia tín dụng và tuân thủ quy trình nội bộ của ngân hàng.'
            )
            disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            disclaimer_run = disclaimer.runs[0]
            disclaimer_run.font.size = Pt(10)
            disclaimer_run.font.italic = True
            disclaimer_run.font.color.rgb = RGBColor(150, 150, 150)

            # Lưu file
            self.doc.save(output_path)
            return output_path

        except Exception as e:
            raise ValueError(f"Lỗi khi tạo báo cáo survival analysis: {str(e)}")


# Instance global
report_generator = ReportGenerator()
